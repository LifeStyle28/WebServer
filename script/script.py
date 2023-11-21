import docx
import sys
import re
import json
import os
import logging

from docx import Document
from num2words import num2words

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell
from docx.enum.table import WD_ROW_HEIGHT_RULE

import datetime
from dateutil.relativedelta import *

# argv[1] - строка шаблона json
# argv[2] - путь для сохранения файла
# argv[3] - номера документов
# argv[4] - срок договора
# argv[5] - процентная ставка

def convert(str):
        return num2words(str, lang='ru')

def space_num(num):
    return format(num, ',d').replace(',',' ')

def paragraph_replace_text(paragraph, regex, replace_str):
    while True:
        text = paragraph.text
        match = regex.search(text)
        if not match:
            break

        runs = iter(paragraph.runs)
        start, end = match.start(), match.end()

        for run in runs:
            run_len = len(run.text)
            if start < run_len:
                break
            start, end = start - run_len, end - run_len

        run_text = run.text
        run_len = len(run_text)
        run.text = "%s%s%s" % (run_text[:start], replace_str, run_text[end:])
        end -= run_len
        for run in runs:
            if end <= 0:
                break
            run_text = run.text
            run_len = len(run_text)
            run.text = run_text[end:]
            end -= run_len
    return paragraph

def fill_cell_table(table, column_num, text, row_set = 1):
        for cell_idx, value in enumerate(table.columns[column_num].cells): # Дата погашения
                                if not(cell_idx > row_set and (cell_idx < int(sys.argv[4]) * 12 + 1 + row_set)):
                                        continue
                                value.paragraphs[0].runs[0].text = text

def fill_last_payment(table, column_num, payment, sum_payments, months, currency = ''):
        table.columns[column_num].cells[-2].paragraphs[0].runs[0].text = space_num(sum_payments - payment * (months - 1)) + currency

def set_cell_border(cell: _Cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def copy_paragraph_style(sample_par, par_to_edit):
        run = sample_par.runs[0]
        output_run = par_to_edit.runs[0]
        # Run's bold data
        output_run.bold = run.bold
        # Run's italic data
        output_run.italic = run.italic
        # Run's underline data
        output_run.underline = run.underline
        # Run's color data
        output_run.font.color.rgb = run.font.color.rgb
        # Run's font data
        output_run.style.name = run.style.name
        output_run.font.name = run.font.name
        output_run.font.size = run.font.size
        # Paragraph's alignment data
        par_to_edit.paragraph_format.alignment = sample_par.paragraph_format.alignment

def create_row_with_borders(table, insert_position):
        table.add_row() # создаем новую строку в конце
        new_row = table.rows[-1] # берем последнюю строку
        table.rows[insert_position]._tr.addnext(new_row._tr) # переносим её на нужную позицию

        for index, cell in enumerate(table.rows[insert_position + 1].cells):
                cell.paragraphs[0].add_run('-')
                copy_paragraph_style(table.cell(insert_position, index).paragraphs[0], cell.paragraphs[0])
                set_cell_border(
                        cell,
                        top={"sz": 4, "val": "single", "color": "#000000"},
                        bottom={"sz": 4, "val": "single", "color": "#000000"},
                        start={"sz": 4, "val": "single", "color": "#000000"},
                        end={"sz": 4, "val": "single", "color": "#000000"}
                )

def create_rows(table, insert_position, count):
        for i in range(count):
                create_row_with_borders(table, insert_position)
                table.rows[insert_position + 1].cells[0].paragraphs[0].runs[0].text = str(count - i + 1)

def fill_cell_date(table, column_num, use_date, row_set = 1):
        create_rows(table, row_set + 1, 12 * int(sys.argv[4]) - 1)
        for cell_idx, value in enumerate(table.columns[column_num].cells): # Дата погашения
                if not(cell_idx > row_set and (cell_idx <= (12 * int(sys.argv[4])) + row_set)):
                        continue
                use_date = use_date + relativedelta(months=+1)
                value.paragraphs[0].runs[0].text = use_date.strftime("%d.%m.%Y")

def fill_with_minuses(table, column_num, row_set = 1):
        for cell_idx, value in enumerate(table.columns[column_num].cells):
                if cell_idx > row_set + 1 and (cell_idx <= (12 * int(sys.argv[4]))+ row_set):
                        value.paragraphs[0].runs[0].text = '-'

def make_tables(document, type, tag_dict):
        logging.debug(f'table type with num={type}')
        if type <= 4:
                make_table_1(document, tag_dict)
        elif type <= 9:
                make_table_2(document, tag_dict, ' рублей')
        elif type <= 10:
                make_table_3(document, tag_dict, ' долларов США')
        elif type <= 11:
                make_table_3(document, tag_dict, ' Евро')
        elif type == 12:
                make_table_2(document, tag_dict, ' USDT TRC-20')

def make_table_1(document, tag_dict):
        years = int(sys.argv[4])
        months = years * 12
        logging.debug('make_table_1')
        date = datetime.datetime.strptime(tag_dict['@<DATE>@'], '%d.%m.%Y')
        loan_sum = int(tag_dict['@<SUMM_NUMBER>@'].replace(' ', ''))

        sum_payments = (loan_sum * int(tag_dict['@<PERCENT_NUMBER>@']) * years) // 100
        payment = (sum_payments // (months * 100)) * 100

        hold_sum = payment * 13 // 100
        if (payment * 13) % 100 > 0:
                hold_sum = hold_sum + 1

        result_sum = payment - hold_sum

        # проверка на нужную таблицу
        for table in document.tables:
                if table.rows[0].cells[0].text  == '№ п/п':
                        fill_cell_date(table, 1, date) # дата погашения
                        fill_cell_table(table, 2, space_num(payment)) # сумма платежа
                        fill_last_payment(table, 2, payment, sum_payments, months) # последний платеж
                        fill_cell_table(table, 3, space_num(hold_sum)) # удерживаемая сумма налога
                        fill_cell_table(table, 4, space_num(result_sum)) # итоговая сумма выплаты Займодавцу
                        fill_with_minuses(table, 5) # сумма основного долга, выставление минусов
                        table.columns[5].cells[-2].paragraphs[0].runs[0].text = space_num(loan_sum) + ' (' + str(convert(loan_sum)) + ') рублей' # сумма основного долга

                        # Заполнение ИТОГО
                        table.rows[-1].cells[2].paragraphs[0].runs[0].text = space_num(sum_payments) + ' (' + str(convert(sum_payments)) + ') рублей' # сумма платежа
                        table.rows[-1].cells[3].paragraphs[0].runs[0].text = space_num(hold_sum * months) + ' (' + str(convert(hold_sum * months)) + ') рублей' # удерживаемая сумма налога
                        table.rows[-1].cells[4].paragraphs[0].runs[0].text = space_num(result_sum * months) + ' (' + str(convert(result_sum * months)) + ') рублей' # утоговая сумма выплаты Займодавцу
                        table.rows[-1].cells[5].paragraphs[0].runs[0].text = space_num(loan_sum) + ' (' + str(convert(loan_sum)) + ') рублей' # сумма основного долга

def make_table_2(document, tag_dict, currency):
        logging.debug('make_table_2')
        years = int(sys.argv[4])
        months = years * 12
        date = datetime.datetime.strptime(tag_dict['@<DATE>@'], '%d.%m.%Y')
        loan_sum = int(tag_dict['@<SUMM_NUMBER>@'].replace(' ', ''))

        sum_payments = (loan_sum * int(tag_dict['@<PERCENT_NUMBER>@']) * years) // 100
        payment = (sum_payments // (months * 100)) * 100

        # проверка на нужную таблицу
        for table in document.tables:
                if table.rows[0].cells[0].text  == '№ п/п':
                        table.rows[0].cells[2].paragraphs[0].runs[0].text = 'Сумма процентов: ' + str(int(tag_dict['@<PERCENT_NUMBER>@'])) + '%'

                        fill_cell_date(table, 1, date) # дата погашения
                        fill_cell_table(table, 2, space_num(payment) + currency) # сумма платежа
                        fill_last_payment(table, 2, payment, sum_payments, months, currency) # последний платеж

                        table.columns[3].cells[1 + months].paragraphs[0].runs[0].text = space_num(loan_sum) + currency # сумма основного долга

                        # Заполнение ИТОГО
                        table.rows[-1].cells[2].paragraphs[0].runs[0].text = space_num(sum_payments) + ' (' + str(convert(sum_payments)) + ')' + currency # сумма платежа
                        if (currency != ' USDT TRC-20'):
                                table.rows[-1].cells[3].paragraphs[0].runs[0].text = space_num(loan_sum) + ' (' + str(convert(loan_sum)) + ')' + currency # сумма основного долга

def make_table_3(document, tag_dict, currency):
        logging.debug('make_table_3')
        years = int(sys.argv[4])
        months = years * 12
        date = datetime.datetime.strptime(tag_dict['@<DATE>@'], '%d.%m.%Y')
        loan_sum = int(tag_dict['@<SUMM_NUMBER>@'].replace(' ', ''))

        sum_payments = (loan_sum * int(tag_dict['@<PERCENT_NUMBER>@']) * years) // 100
        payment = (sum_payments // (months * 100)) * 100

        # проверка на нужную таблицу
        for table in document.tables:
                if table.rows[0].cells[0].text == '№':

                        fill_cell_date(table, 1, date, 0) # дата погашения
                        fill_cell_table(table, 2, space_num(payment) + currency + ' в российских рублях по курсу ЦБ РФ на день выплаты', 0) # сумма платежа
                        fill_last_payment(table, 2, payment, sum_payments, months, currency) # последний платеж

                        # заполнение ИТОГО
                        table.rows[months].cells[2].paragraphs[0].runs[0].text = 'Возвращается Сумма займа в размере ' + space_num(sum_payments) + ' (' + str(convert(sum_payments)) + ') ' + currency # сумма платежа

def get_docs_nums(data): # получим номера нужных документов
        doc_nums = []
        for num in data.split():
                doc_nums.append(int(num))
        return doc_nums

def tables_replace_text(document, regex, new_value):
        for table in document.tables:
                for row in table.rows:
                        for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                        paragraph_replace_text(paragraph, regex, new_value)
def get_name_by_num(num):
    file_names = {1: 'Рустонн + физ. лицо (Руб)',
                  5: 'Рустонн + юр. лицо (Руб)',
                  9: 'Пугачев Т.В. + физ. лицо (Руб)',
                  10: 'Пугачев Т.В. + физ. лицо ($)',
                  11: 'Пугачев Т.В. + физ. лицо (Euro)',
                  12: 'Пугачев Т.В. + физ. лицо (USDT)',
                  13: 'Поручительство Пугачев Т.В. (Руб)',
                  17: 'Поручительство Рустонн (Руб)',
                  18: 'Поручительство Рустонн ($)',
                  19: 'Поручительство Рустонн (Euro)',
                  20: 'Поручительство Рустонн (USDT)',
                  21: 'Расписка (Руб)',
                  22: 'Расписка ($)',
                  23: 'Расписка (Euro)'}
    return file_names[num]

def correct_values(tag_dict):
        if '@<SUMM_NUMBER>@' in tag_dict:
                tag_dict['@<SUMM_NUMBER>@'] = space_num(int(tag_dict['@<SUMM_NUMBER>@'])) # если это сумма договора, то разделяем пробелами

def replace_tags(document, tag_dict):
        for tag, value in tag_dict.items(): # бежим по тэгам
                regex = re.compile(tag)
                for paragraph in document.paragraphs:
                        paragraph_replace_text(paragraph, regex, value)
                tables_replace_text(document, regex, value)

def make_fio_short(fio_full, char = ''):
        try:
                split_fio = fio_full.split()
                result = ''
                for i, word in enumerate(split_fio):
                        if 0 == i and '.' == char: # нужны фамилия и инициалы через точку
                                result = word + ' '
                        else:
                                result += word[0] + char
                return result
        except Exception as e: # в случае любой ошибки возвращаем первоначальные ФИО
                return fio_full

def split_passport(seria_num):
        if len(seria_num.split()) == 1:
                return seria_num[0:4], seria_num[4:]
        return seria_num.split()[0], seria_num.split()[1]

def make_contract_num(tag_dict):
        date = datetime.datetime.strptime(tag_dict['@<DATE>@'], '%d.%m.%Y')
        datestr = date.strftime('%d%m%y')
        return datestr + make_fio_short(tag_dict['@<FIO_FULL>@'])

def add_extra_tags(tag_dict):
        if '@<SUMM_NUMBER>@' in tag_dict:
                tag_dict['@<SUMM_TEXT>@'] = str(convert(int(tag_dict['@<SUMM_NUMBER>@'])))
        if '@<FIO_FULL>@' in tag_dict:
                tag_dict['@<FIO_SHORT>@'] = make_fio_short(tag_dict['@<FIO_FULL>@'], '.')
        if '@<PASSPORT_SERIA_NUM>@' in tag_dict:
                tag_dict['@<PASSPORT_SERIA>@'], tag_dict['@<PASSPORT_NUM>@'] = split_passport(tag_dict['@<PASSPORT_SERIA_NUM>@'])
        tag_dict['@<GUARANTEE_NUM>@'] = str(make_contract_num(tag_dict))
        tag_dict['@<CONTRACT_NUM>@'] = str(make_contract_num(tag_dict))
        date = datetime.datetime.strptime(tag_dict['@<DATE>@'], '%d.%m.%Y') + relativedelta(months=+(int(sys.argv[4]) * 12))
        tag_dict['@<CONTRACT_TERM>@'] = str(date.strftime("%d.%m.%Y"))
        tag_dict['@<PERCENT_NUMBER>@'] = str(sys.argv[5])
        if '@<PERCENT_NUMBER>@' in tag_dict:
                tag_dict['@<PERCENT_TEXT>@'] = str(convert(int(tag_dict['@<PERCENT_NUMBER>@'])))
        if int(sys.argv[4]) == 2:
                months_text = ' месяца'
        else:
                months_text = ' месяцев'
        tag_dict['@<MONTHS>@'] = str(int(sys.argv[4]) * 12) + months_text

def make_fio_date(tag_dict):
        return tag_dict['@<DATE>@'] + ' ' + tag_dict['@<FIO_SHORT>@']

def human_format(num):
    num = float('{:.3g}'.format(num))
    magnitude = 0
    while abs(num) >= 1000:
        magnitude += 1
        num /= 1000.0
    return '{}{}'.format('{:f}'.format(num).rstrip('0').rstrip('.'), ['', 'к', 'кк', 'ккк', 'кккк'][magnitude])

def make_zip_name(date, sum, currency_char, fio_short):
        fio_short = fio_short.replace('.', '')
        fio_short = fio_short.replace(' ', '')
        result = date.replace('.', '') + '_' + human_format(int(sum.replace(' ', ''))) + '_' + currency_char + '_' + fio_short

        return result;

def get_currency_char(contract_num):
        if contract_num == 10:
                return '$'
        elif contract_num == 11:
                return '€'
        elif contract_num == 12:
                return '₮'
        else:
                return '₽'

def create_extra_field_file(tag_dict):
        if len(tag_dict['@<EXTRA_FIELD>@']) != 0: # если есть доп.поле, то вывести в отдельный файл
                os.system('echo "' + tag_dict['@<EXTRA_FIELD>@'] + '" > ' + str(sys.argv[2]) + 'Примечания.txt')

def make_docs():
        data = json.loads(str(sys.argv[1])) # загружаем шаблон
        tag_dict = {} # создаем словарь соответствия тэгов и значений
        for entry in data['from_tag_values']: # заполняем словарь тэгов
                tag_dict[entry['tag']] = str(entry['value'])
        add_extra_tags(tag_dict) # добавить новые необходимые теги, такие как текстовые представления чисел
        correct_values(tag_dict) # подмениваем значения по необходимости
        logging.debug(f'tag_dict = : {tag_dict}')

        doc_nums = get_docs_nums(str(sys.argv[3])) # получаем номера документов
        for doc_num in doc_nums: # бежим по номерам документов
                document = docx.Document(DOCX_TEMPLATES_PATH + str(doc_num) + '.docx') # открываем шаблон документа
                replace_tags(document, tag_dict) # заменяем тэги на нужные значения
                make_tables(document, doc_num, tag_dict) # заполняем таблицы
                save_name = str(sys.argv[2] + get_name_by_num(doc_num) + ' от ' + make_fio_date(tag_dict) + ".docx")
                logging.debug(f"Name to save is {save_name}")
                document.save(save_name)
        create_extra_field_file(tag_dict)
        command = 'cd ' + str(sys.argv[2]) + ' && 7z a ' + make_zip_name(tag_dict['@<DATE>@'], tag_dict['@<SUMM_NUMBER>@'], get_currency_char(doc_nums[0]), tag_dict['@<FIO_SHORT>@']) + '.7z ' + './'
        logging.debug(command)
        os.system(command)

if __name__ == "__main__":
        LOG_FULL_NAME = '/app/script/python_script.log'
        DOCX_TEMPLATES_PATH = '/app/templates/docx/'

        logging.basicConfig(level=logging.DEBUG, filename=LOG_FULL_NAME) # запись логов уровня дебаг в файл LOG_FULL_NAME
        logging.debug(f'Args length: {len(sys.argv)}')
        logging.debug(f'Args = {str(sys.argv)}')

        make_docs()
